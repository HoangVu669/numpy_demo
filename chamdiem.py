# chạy 2 lệnh này để inporrt thư viện 
# pip install python-docx
# pip install numpy

import os
import re
import numpy as np
from docx import Document

def load_student_answers(folder_path):
    student_answers = {}
    try:
        for file_name in os.listdir(folder_path):
            if file_name.endswith('.docx'):
                file_path = os.path.join(folder_path, file_name)
                doc = Document(file_path)
                student_id = file_name.split('.')[0]  # Giả sử tên tệp là ID sinh viên
                answers = []

                # Mẫu regex để xác định các dòng câu hỏi
                question_pattern = re.compile(r'^(Câu )?\d+:.*$|^(C)?\d+:.*$')

                for p in doc.paragraphs:
                    line = p.text.strip()
                    if question_pattern.match(line):
                        parts = line.split(':')
                        if len(parts) == 2:
                            try:
                                question_num = int(re.search(r'\d+', parts[0]).group())  # Lấy số câu hỏi từ phần đầu của dòng
                                answer = parts[1].strip().upper()
                                answers.append((question_num, answer))
                            except AttributeError:
                                continue

                # Sắp xếp câu trả lời theo số câu hỏi
                answers.sort(key=lambda x: x[0])

                # Lấy chỉ câu trả lời từ danh sách đã sắp xếp
                answers = [answer[1] for answer in answers]

                # if len(answers) >= 4:  # Kiểm tra xem tệp có chứa đủ thông tin sinh viên và câu trả lời không
                #     student_answers[student_id] = answers
                # else:
                #     print(f"Bỏ qua tệp {file_name} vì không chứa đủ thông tin sinh viên hoặc câu trả lời.")
    except FileNotFoundError:
        print("Thư mục không được tìm thấy. Vui lòng kiểm tra đường dẫn thư mục.")
    except Exception as e:
        print("Đã xảy ra lỗi:", e)
    return student_answers




def load_answer_key(answer_key_path):
    try:
        doc = Document(answer_key_path)
        answer_key = [p.text.strip().split(':')[-1].strip().upper() for p in doc.paragraphs if p.text.strip()]
    except FileNotFoundError:
        print("Tệp đáp án không được tìm thấy. Vui lòng kiểm tra đường dẫn.")
        return None
    except Exception as e:
        print("Đã xảy ra lỗi khi tải đáp án:", e)
        return None
    return answer_key

def grade_exam(student_answers, answer_key):
    scores = {}
    try:
        for student_id, answers in student_answers.items():
            num_correct = np.sum([1 if answer == key else 0 for answer, key in zip(answers, answer_key)])
            student_score = (10 / len(answer_key)) * num_correct
            scores[student_id] = student_score
    except Exception as e:
        print("Đã xảy ra lỗi khi chấm điểm:", e)
    return scores

def generate_report(scores, output_file):
    try:
        with open(output_file, 'w', encoding='utf-8') as file:
            for student_id, score in scores.items():
                file.write(f"Mã sinh viên: {student_id}, Điểm: {score}\n")
        print("Báo cáo đã được tạo thành công.")
    except Exception as e:
        print("Đã xảy ra lỗi khi tạo báo cáo:", e)

if __name__ == "__main__":
    folder_path = "lop20cn2"  # Đường dẫn tới thư mục chứa các tệp Word của học sinh
    student_answers = load_student_answers(folder_path)
    for student_id, answers in student_answers.items():
        print(f"Mã sinh viên: {student_id}, Câu trả lời: {answers}")

    answer_key_path = r"C:\Users\hoang\Desktop\project-python\dapan.docx"
    answer_key = load_answer_key(answer_key_path)
    if answer_key:
        scores = grade_exam(student_answers, answer_key)
        output_file = "ket_qua_thi.txt"  # Tên của tệp báo cáo kết quả
        generate_report(scores, output_file)
